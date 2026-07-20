"""Document difference detection: compare a .docx against specimen documents."""

import difflib
import io
import re

from docx import Document
from docx.shared import Pt, RGBColor

_QUOTE_MAP = str.maketrans({
    "\u2018": "'", "\u2019": "'",  # curly single quotes
    "\u201C": '"', "\u201D": '"',  # curly double quotes
    "\u2013": "-", "\u2014": "-",  # en/em dashes
    "\u00A0": " ",                  # non-breaking space
})


def _normalize(text):
    """Normalize unicode punctuation for comparison purposes."""
    return text.translate(_QUOTE_MAP)


def extract_text_blocks(docx_bytes):
    """Extract paragraph texts from a .docx file as a list of strings."""
    doc = Document(io.BytesIO(docx_bytes))
    blocks = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            blocks.append(text)
    return blocks


def compute_similarity(blocks_a, blocks_b):
    """Compute similarity ratio between two documents' text blocks (0.0–1.0)."""
    text_a = _normalize("\n".join(blocks_a))
    text_b = _normalize("\n".join(blocks_b))
    return difflib.SequenceMatcher(None, text_a, text_b).ratio()


def rank_specimens(upload_bytes, specimens):
    """Rank specimen documents by similarity to the uploaded document.

    Args:
        upload_bytes: raw .docx bytes of the uploaded document
        specimens: list of dicts with 'id', 'name', 'docx_data' (bytes)

    Returns:
        list of dicts: [{'id', 'name', 'similarity'}, ...] sorted desc by similarity
    """
    upload_blocks = extract_text_blocks(upload_bytes)
    results = []
    for spec in specimens:
        spec_blocks = extract_text_blocks(spec["docx_data"])
        sim = compute_similarity(upload_blocks, spec_blocks)
        results.append({
            "id": str(spec["id"]),
            "name": spec["name"],
            "similarity": round(sim * 100, 1),
        })
    results.sort(key=lambda x: x["similarity"], reverse=True)
    return results


def _map_replace_block(spec_paras, upload_paras, j_start, diff_map):
    """Within a replace block, match each upload paragraph to its specimen counterpart.

    Equal-length blocks are paired by index (positional alignment from the outer
    SequenceMatcher). Unequal blocks use greedy similarity matching, then
    force-pair any remaining unmatched paragraphs in order so short placeholder
    fill-ins are not mislabeled as Added/Removed.
    """
    n_spec = len(spec_paras)
    n_up = len(upload_paras)
    pairs = {}  # up_idx -> spec_idx

    if n_spec == n_up:
        for i in range(n_up):
            pairs[i] = i
    else:
        used_spec = set()
        for up_idx, up_text in enumerate(upload_paras):
            best_sim = -1
            best_spec_idx = None
            for s_idx, spec_text in enumerate(spec_paras):
                if s_idx in used_spec:
                    continue
                sim = difflib.SequenceMatcher(
                    None, _normalize(spec_text), _normalize(up_text)
                ).ratio()
                if sim > best_sim:
                    best_sim = sim
                    best_spec_idx = s_idx
            if best_spec_idx is not None and best_sim > 0.4:
                used_spec.add(best_spec_idx)
                pairs[up_idx] = best_spec_idx

        leftover_up = [i for i in range(n_up) if i not in pairs]
        leftover_spec = [i for i in range(n_spec) if i not in used_spec]
        for up_idx, s_idx in zip(leftover_up, leftover_spec):
            pairs[up_idx] = s_idx

    used_spec = set(pairs.values())
    for up_idx, up_text in enumerate(upload_paras):
        abs_idx = j_start + up_idx
        if up_idx in pairs:
            word_diff = _word_level_diff(spec_paras[pairs[up_idx]], up_text)
            if word_diff:
                diff_map[abs_idx] = word_diff
        else:
            diff_map[abs_idx] = "Added — not in specimen"

    unmatched = [spec_paras[i] for i in range(n_spec) if i not in used_spec]
    if unmatched:
        removed = " | ".join(unmatched)
        if len(removed) > 300:
            removed = removed[:300] + "..."
        note = f"Removed from specimen: \"{removed}\""
        first_idx = j_start
        if first_idx in diff_map:
            diff_map[first_idx] += f" | {note}"
        else:
            diff_map[first_idx] = note


def _word_level_diff(specimen_para, upload_para):
    """Produce a concise description of word-level changes between two paragraphs.

    Returns None if paragraphs are identical after normalization.
    """
    norm_spec = _normalize(specimen_para)
    norm_up = _normalize(upload_para)

    if norm_spec == norm_up:
        return None

    spec_words = norm_spec.split()
    up_words = norm_up.split()

    matcher = difflib.SequenceMatcher(None, spec_words, up_words)
    changes = []

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            continue
        old = " ".join(spec_words[i1:i2])
        new = " ".join(up_words[j1:j2])
        if tag == "replace":
            changes.append(f"\"{old}\" \u2192 \"{new}\"")
        elif tag == "insert":
            changes.append(f"Added: \"{new}\"")
        elif tag == "delete":
            changes.append(f"Removed: \"{old}\"")

    if not changes:
        return None

    if len(changes) <= 5:
        return " | ".join(changes)

    return " | ".join(changes[:5]) + f" | ... and {len(changes) - 5} more changes"


def build_diff_docx(upload_bytes, specimen_bytes):
    """Build a .docx with comments highlighting differences from the specimen.

    Returns a BytesIO containing the annotated .docx.
    """
    upload_doc = Document(io.BytesIO(upload_bytes))
    specimen_blocks = extract_text_blocks(specimen_bytes)

    upload_blocks = []
    for para in upload_doc.paragraphs:
        text = para.text.strip()
        if text:
            upload_blocks.append(text)

    norm_spec = [_normalize(b) for b in specimen_blocks]
    norm_up = [_normalize(b) for b in upload_blocks]
    matcher = difflib.SequenceMatcher(None, norm_spec, norm_up)
    opcodes = matcher.get_opcodes()

    diff_map = {}
    for tag, i1, i2, j1, j2 in opcodes:
        if tag == "equal":
            continue
        elif tag == "replace":
            spec_range = specimen_blocks[i1:i2]
            up_range = upload_blocks[j1:j2]
            _map_replace_block(spec_range, up_range, j1, diff_map)
        elif tag == "insert":
            for idx in range(j1, j2):
                diff_map[idx] = "Added — not in specimen"
        elif tag == "delete":
            removed = " | ".join(specimen_blocks[i1:i2])
            if len(removed) > 300:
                removed = removed[:300] + "..."
            note = f"Removed from specimen: \"{removed}\""
            if j1 in diff_map:
                diff_map[j1] += f" | {note}"
            else:
                nearest = min(j1, len(upload_blocks) - 1)
                diff_map[nearest] = note

    out_doc = Document()
    style = out_doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    upload_block_idx = 0
    for para in upload_doc.paragraphs:
        text = para.text.strip()
        if not text:
            out_doc.add_paragraph("")
            continue

        new_para = out_doc.add_paragraph()
        _copy_para_format(para, new_para)
        runs = []
        for run in para.runs:
            new_run = new_para.add_run(run.text)
            _copy_run_format(run, new_run)
            runs.append(new_run)

        if upload_block_idx in diff_map and runs:
            comment_text = diff_map[upload_block_idx]
            if len(comment_text) > 1000:
                comment_text = comment_text[:1000] + "..."
            out_doc.add_comment(runs=runs, text=comment_text, author="Difference")

        upload_block_idx += 1

    buf = io.BytesIO()
    out_doc.save(buf)
    buf.seek(0)
    return buf


def _copy_para_format(src, dst):
    """Copy basic paragraph formatting."""
    if src.paragraph_format.alignment is not None:
        dst.paragraph_format.alignment = src.paragraph_format.alignment


def _copy_run_format(src, dst):
    """Copy basic run formatting."""
    dst.bold = src.bold
    dst.italic = src.italic
    dst.underline = src.underline
    if src.font.size:
        dst.font.size = src.font.size
    if src.font.name:
        dst.font.name = src.font.name
