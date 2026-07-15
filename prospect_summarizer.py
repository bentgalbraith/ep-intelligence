"""Prospect EP document extraction: OCR via Google Document AI, structured extraction via Grok."""

import io
import json
import logging
import os
import re
import time
import traceback
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from google.cloud import documentai_v1 as documentai
from pypdf import PdfReader, PdfWriter

from doc_separator import _get_documentai_client, _page_text, PAGES_PER_CHUNK

log = logging.getLogger("prospect_summarizer")
log.setLevel(logging.DEBUG)
if not log.handlers:
    _h = logging.StreamHandler()
    _h.setFormatter(logging.Formatter("[prospect_summarizer] %(message)s"))
    log.addHandler(_h)

SCHEMA_PATH = Path(__file__).parent / "prospect_schema.json"
if SCHEMA_PATH.exists():
    with open(SCHEMA_PATH) as f:
        PROSPECT_SCHEMA = json.load(f)
else:
    PROSPECT_SCHEMA = {"sections": []}

EXTRACTION_PROMPT_TEMPLATE = """\
You are a legal assistant specializing in estate planning.{firm_context}

You will receive OCR text from a prospective client's existing estate planning \
documents. Your job is to extract structured data from the documents according to \
a schema I will provide.

For EVERY field you can answer, return:
- "value": a concise, form-ready answer. For short-answer fields (names, dates, \
jurisdictions), the exact text from the document is fine. For longer narrative fields, \
synthesize the relevant information into a clean summary.
- "quote": the exact phrase(s) from the OCR text that support your answer. \
Copy the words verbatim — preserve the original spelling and punctuation even if \
the OCR introduced errors. If multiple parts of the text are relevant, join them \
with " ... ". Every non-null value MUST have an accompanying quote.

If a field cannot be answered from the documents, set both value and quote to null. \
Do not guess or fabricate.

Important rules:
- Focus on the TRUST DOCUMENTS. Ignore ancillary documents (powers of attorney, \
healthcare surrogates, living wills, pour-over wills) for the purposes of this \
extraction, unless they contain information relevant to the trust structure.
- Synthesize across all documents. Do not produce per-document answers.
- For the "Other Notes" field, include any notable provisions, unusual clauses, \
or important context that does not fit in the other fields but that an attorney \
should know before a meeting.

Each field in the schema has a "type" property. Handle each type as follows:
- "text": set "value" to a concise synthesized answer. Capitalize the first letter.
- "choice": set "value" to exactly one of the strings listed in the field's \
"options" array. Use the exact spelling and capitalization from the options.

Return ONLY valid JSON with this structure (no markdown, no commentary):
{{
  "sections": {{
    "<section_id>": {{
      "fields": {{
        "<field_id>": {{"value": "...", "quote": "..."}},
        ...
      }}
    }}
  }}
}}

Here is the schema:
"""


def _build_extraction_prompt(firm_context=""):
    ctx = ""
    if firm_context:
        ctx = f" {firm_context}"
    return EXTRACTION_PROMPT_TEMPLATE.format(firm_context=ctx)


def _ocr_pdf(pdf_content, firm_id=None):
    """OCR a single PDF via Document AI; return concatenated page-labeled text."""
    from ai_logger import log_ai_call

    reader = PdfReader(io.BytesIO(pdf_content))
    total = len(reader.pages)
    ocr_start = time.time()
    client, processor_name = _get_documentai_client()

    texts = {}
    for start in range(0, total, PAGES_PER_CHUNK):
        end = min(start + PAGES_PER_CHUNK, total)
        writer = PdfWriter()
        for i in range(start, end):
            writer.add_page(reader.pages[i])

        buf = io.BytesIO()
        writer.write(buf)

        try:
            result = client.process_document(
                request=documentai.ProcessRequest(
                    name=processor_name,
                    raw_document=documentai.RawDocument(
                        content=buf.getvalue(), mime_type="application/pdf"
                    ),
                )
            )
        except Exception:
            log_ai_call(
                provider="google_documentai", tool="prospect_summarizer_ocr", status="error",
                pages_processed=total,
                execution_ms=int((time.time() - ocr_start) * 1000),
                notes=traceback.format_exc(),
                firm_id=firm_id,
            )
            raise

        for local_idx, page in enumerate(result.document.pages):
            texts[start + local_idx + 1] = _page_text(result.document, page)

    log_ai_call(
        provider="google_documentai", tool="prospect_summarizer_ocr", status="success",
        pages_processed=total,
        execution_ms=int((time.time() - ocr_start) * 1000),
        firm_id=firm_id,
    )

    return texts, total


def extract_prospect_documents(pdf_contents, client, notes="", model=None,
                                firm_id=None, firm_config=None):
    """OCR multiple PDFs, then ask OpenAI to extract structured data.

    Returns (extraction_dict, ocr_text).
    """
    total_start = time.time()
    model = model or os.environ.get("PROSPECT_SUMMARIZER_MODEL", "gpt-5.6-terra")
    firm_config = firm_config or {}
    log.info("Starting: model=%s, %d PDF(s)", model, len(pdf_contents))

    all_text = ""
    for idx, pdf_bytes in enumerate(pdf_contents, 1):
        log.info("OCR-ing PDF %d (%d bytes)", idx, len(pdf_bytes))
        page_texts, total_pages = _ocr_pdf(pdf_bytes, firm_id=firm_id)
        for pn in range(1, total_pages + 1):
            txt = page_texts.get(pn, "").strip()
            all_text += f"\n--- PDF {idx}, PAGE {pn} ---\n{txt}\n"

    total_chars = len(all_text)
    log.info("OCR complete: %d total chars across all PDFs", total_chars)

    user_content = all_text
    if notes.strip():
        user_content += f"\n\n--- ADDITIONAL CONTEXT ---\n{notes.strip()}\n"

    prospect_schema = firm_config.get("prospect_schema") or PROSPECT_SCHEMA
    schema_text = json.dumps(prospect_schema.get("sections", []), indent=2)
    firm_context = firm_config.get("firm_context", "")
    system_content = _build_extraction_prompt(firm_context) + schema_text

    from ai_logger import log_ai_call, extract_xai_usage
    from doc_separator import _extract_json

    call_start = time.time()
    try:
        resp = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": system_content},
                {"role": "user", "content": user_content},
            ],
        )
    except Exception:
        log_ai_call(
            provider="openai", model=model, tool="prospect_summarizer", status="error",
            execution_ms=int((time.time() - call_start) * 1000),
            notes=traceback.format_exc(),
            firm_id=firm_id,
        )
        raise
    call_elapsed = time.time() - call_start

    log_ai_call(
        provider="openai", model=model, tool="prospect_summarizer", status="success",
        execution_ms=int(call_elapsed * 1000),
        firm_id=firm_id,
        **extract_xai_usage(resp),
    )

    usage = resp.usage
    log.info(
        "OpenAI: %.1fs, %s/%s tokens (prompt/completion)",
        call_elapsed,
        getattr(usage, "prompt_tokens", "?"),
        getattr(usage, "completion_tokens", "?"),
    )

    raw = resp.choices[0].message.content.strip()
    if raw.startswith("```"):
        raw = raw.split("\n", 1)[1].rsplit("```", 1)[0].strip()

    extraction = _extract_json(raw)

    total_elapsed = time.time() - total_start
    log.info("Complete: %.1fs total", total_elapsed)
    return extraction, all_text


def _join_names(names):
    """Join a list of names with commas and 'and' before the last."""
    if not names:
        return "no documents"
    if len(names) == 1:
        return names[0]
    return ", ".join(names[:-1]) + ", and " + names[-1]


_SANS = "Calibri"


def _sans_run(run, size_pt=None, bold=None, italic=None):
    run.font.name = _SANS
    if size_pt is not None:
        run.font.size = size_pt
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def build_summary_docx(sections_data, schema, filenames=None):
    """Build a Word doc from the structured extraction. Returns BytesIO."""
    from datetime import datetime
    from docx.shared import RGBColor

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = _SANS
    style.font.size = Pt(11)
    for s in doc.sections:
        s.top_margin = Inches(1)
        s.bottom_margin = Inches(1)
        s.left_margin = Inches(1)
        s.right_margin = Inches(1)

    heading = doc.add_paragraph()
    run = heading.add_run("Prospect EP Summary")
    run.font.name = _SANS
    run.font.size = Pt(14)
    run.bold = True
    heading.paragraph_format.space_after = Pt(2)

    from zoneinfo import ZoneInfo
    now = datetime.now(ZoneInfo("America/New_York"))
    date_str = now.strftime("%B %d, %Y")
    time_str = now.strftime("%I:%M %p").lstrip("0")
    subtitle_text = f"Generated by AI on {date_str} at {time_str}."
    if filenames:
        subtitle_text = f"Generated by AI on {date_str} at {time_str} using the following documents: {_join_names(filenames)}."

    subtitle = doc.add_paragraph()
    sub_run = subtitle.add_run(subtitle_text)
    sub_run.font.name = _SANS
    sub_run.font.size = Pt(9)
    sub_run.italic = True
    sub_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    subtitle.paragraph_format.space_after = Pt(12)

    sections_data = sections_data or {}

    for sec_idx, sec in enumerate(schema):
        sid = sec["id"]
        sdata = sections_data.get(sid) or {}
        fmap = sdata.get("fields") or {}

        if sec_idx > 0:
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(6)
            spacer.paragraph_format.space_after = Pt(6)

        fields = sec.get("fields") or []
        single_field_matches_section = (
            len(fields) == 1 and fields[0].get("name") == sec["name"]
        )
        if not single_field_matches_section:
            sec_header = doc.add_paragraph()
            sec_header.paragraph_format.space_before = Pt(2)
            sec_header.paragraph_format.space_after = Pt(4)
            r = sec_header.add_run(sec["name"])
            _sans_run(r, size_pt=Pt(11), bold=True)

        for field in sec.get("fields") or []:
            fid = field["id"]
            fd = fmap.get(fid) or {}
            val = fd.get("value")
            quote = fd.get("quote")
            field_name = field.get("name") or fid
            field_type = field.get("type", "text")

            qv = str(val).strip() if val is not None else ""
            qq = str(quote).strip() if quote is not None else ""

            if field_type == "choice":
                _add_choice_line(doc, field_name, qv, field.get("options") or [], qq)
            else:
                _add_field_line(doc, field_name, qv, qq)

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out


def _add_field_line(doc, field_name, value, quote):
    """Single field row: 'Label: answer' with optional Word comment for quote."""
    body = Pt(11)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15

    lab = p.add_run(field_name.strip() + ": ")
    _sans_run(lab, size_pt=body, bold=True)

    value_runs = []

    if value:
        lines = value.replace("\r\n", "\n").split("\n")
        vr = p.add_run(lines[0])
        _sans_run(vr, size_pt=body)
        value_runs.append(vr)
        for extra in lines[1:]:
            cont = doc.add_paragraph(extra)
            cont.paragraph_format.space_after = Pt(0)
            cont.paragraph_format.left_indent = Pt(18)
            for r in cont.runs:
                _sans_run(r, size_pt=body)
                value_runs.append(r)
    else:
        ur = p.add_run("_" * 52)
        _sans_run(ur, size_pt=body)
        value_runs.append(ur)

    if quote and value_runs:
        inner = quote.strip().replace('"', "'")
        comment_text = 'Quote from Document: "' + inner + '"'
        doc.add_comment(runs=value_runs, text=comment_text, author="")


def _add_choice_line(doc, field_name, value, options, quote):
    """Render a choice field with checkbox-style display in the Word doc."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from lxml import etree
    import random

    body = Pt(11)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15

    lab = p.add_run(field_name.strip() + ":  ")
    _sans_run(lab, size_pt=body, bold=True)

    for i, opt in enumerate(options):
        is_selected = (value == opt)

        sdt = OxmlElement("w:sdt")
        sdt_pr = OxmlElement("w:sdtPr")
        sdt.append(sdt_pr)
        sdt_id = OxmlElement("w:id")
        sdt_id.set(qn("w:val"), str(random.randint(100000000, 999999999)))
        sdt_pr.append(sdt_id)
        checkbox = etree.SubElement(sdt_pr, qn("w14:checkbox"))
        checked_el = etree.SubElement(checkbox, qn("w14:checked"))
        checked_el.set(qn("w14:val"), "1" if is_selected else "0")
        checked_state = etree.SubElement(checkbox, qn("w14:checkedState"))
        checked_state.set(qn("w14:val"), "2612")
        checked_state.set(qn("w14:font"), "MS Gothic")
        unchecked_state = etree.SubElement(checkbox, qn("w14:uncheckedState"))
        unchecked_state.set(qn("w14:val"), "2610")
        unchecked_state.set(qn("w14:font"), "MS Gothic")
        sdt.append(OxmlElement("w:sdtEndPr"))
        sdt_content = OxmlElement("w:sdtContent")
        sdt.append(sdt_content)
        run_el = OxmlElement("w:r")
        sdt_content.append(run_el)
        rpr = OxmlElement("w:rPr")
        run_el.append(rpr)
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), "MS Gothic")
        fonts.set(qn("w:eastAsia"), "MS Gothic")
        fonts.set(qn("w:hAnsi"), "MS Gothic")
        fonts.set(qn("w:hint"), "eastAsia")
        rpr.append(fonts)
        t = OxmlElement("w:t")
        t.text = "\u2612" if is_selected else "\u2610"
        run_el.append(t)
        p._element.append(sdt)

        spacer = f" {opt}"
        if i < len(options) - 1:
            spacer += "     "
        r = p.add_run(spacer)
        _sans_run(r, size_pt=body)

    if quote:
        inner = quote.strip().replace('"', "'")
        comment_text = 'Quote from Document: "' + inner + '"'
        all_runs = list(p.runs)
        if all_runs:
            doc.add_comment(runs=all_runs, text=comment_text, author="")
