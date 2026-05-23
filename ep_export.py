"""CSV and Word export for EP extraction (values + transcript quotes)."""

import csv
import random
from io import BytesIO, StringIO

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from lxml import etree

_SANS = "Calibri"


def _sans_run(run, size_pt=None, bold=None, italic=None):
    run.font.name = _SANS
    if size_pt is not None:
        run.font.size = size_pt
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _str_or_empty(v):
    if v is None:
        return ""
    return str(v).strip()


def _add_checkbox_sdt(paragraph, checked=False):
    """Append a real Word w14:checkbox SDT control to a paragraph."""
    sdt = OxmlElement("w:sdt")

    sdt_pr = OxmlElement("w:sdtPr")
    sdt.append(sdt_pr)

    sdt_id = OxmlElement("w:id")
    sdt_id.set(qn("w:val"), str(random.randint(100000000, 999999999)))
    sdt_pr.append(sdt_id)

    checkbox = etree.SubElement(sdt_pr, qn("w14:checkbox"))
    checked_el = etree.SubElement(checkbox, qn("w14:checked"))
    checked_el.set(qn("w14:val"), "1" if checked else "0")
    checked_state = etree.SubElement(checkbox, qn("w14:checkedState"))
    checked_state.set(qn("w14:val"), "2612")
    checked_state.set(qn("w14:font"), "MS Gothic")
    unchecked_state = etree.SubElement(checkbox, qn("w14:uncheckedState"))
    unchecked_state.set(qn("w14:val"), "2610")
    unchecked_state.set(qn("w14:font"), "MS Gothic")

    sdt.append(OxmlElement("w:sdtEndPr"))

    sdt_content = OxmlElement("w:sdtContent")
    sdt.append(sdt_content)

    run = OxmlElement("w:r")
    sdt_content.append(run)

    rpr = OxmlElement("w:rPr")
    run.append(rpr)
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "MS Gothic")
    fonts.set(qn("w:eastAsia"), "MS Gothic")
    fonts.set(qn("w:hAnsi"), "MS Gothic")
    fonts.set(qn("w:hint"), "eastAsia")
    rpr.append(fonts)

    t = OxmlElement("w:t")
    t.text = "\u2612" if checked else "\u2610"
    run.append(t)

    paragraph._element.append(sdt)


def iter_export_rows(sections, schema, repeatable_sections):
    """
    Walk schema order; yield dicts with keys:
    section_id, section_name, entry_label, field_id, field_name, field_description, value, quote
    """
    sections = sections or {}
    repeatable_set = set(repeatable_sections or [])

    for sec in schema:
        sid = sec["id"]
        sname = sec.get("name") or sid
        sdata = sections.get(sid) or {}

        if sid in repeatable_set:
            entries = sdata.get("entries") or []
            if not entries:
                for field in sec.get("fields") or []:
                    yield {
                        "section_id": sid,
                        "section_name": sname,
                        "entry_label": "(none)",
                        "field_id": field["id"],
                        "field_name": field.get("name") or field["id"],
                        "field_description": field.get("description") or "",
                        "value": "",
                        "quote": "",
                    }
                continue
            for ei, ent in enumerate(entries):
                label = f"Entry {ei + 1}"
                fmap = ent.get("fields") or {}
                for field in sec.get("fields") or []:
                    fid = field["id"]
                    fd = fmap.get(fid) or {}
                    val = fd.get("value")
                    quote = fd.get("quote")
                    yield {
                        "section_id": sid,
                        "section_name": sname,
                        "entry_label": label,
                        "field_id": fid,
                        "field_name": field.get("name") or fid,
                        "field_description": field.get("description") or "",
                        "value": _str_or_empty(val) if val is not None else "",
                        "quote": _str_or_empty(quote) if quote is not None else "",
                    }
        else:
            fmap = sdata.get("fields") or {}
            for field in sec.get("fields") or []:
                fid = field["id"]
                fd = fmap.get(fid) or {}
                val = fd.get("value")
                quote = fd.get("quote")
                yield {
                    "section_id": sid,
                    "section_name": sname,
                    "entry_label": "",
                    "field_id": fid,
                    "field_name": field.get("name") or fid,
                    "field_description": field.get("description") or "",
                    "value": _str_or_empty(val) if val is not None else "",
                    "quote": _str_or_empty(quote) if quote is not None else "",
                }


def build_export_csv(sections, schema, repeatable_sections):
    """UTF-8 with BOM for Excel; returns BytesIO."""
    buf = StringIO()
    writer = csv.writer(buf)
    writer.writerow(
        [
            "section_id",
            "section_name",
            "entry",
            "field_id",
            "field_name",
            "field_description",
            "value",
            "quote",
        ]
    )
    for row in iter_export_rows(sections, schema, repeatable_sections):
        writer.writerow(
            [
                row["section_id"],
                row["section_name"],
                row["entry_label"],
                row["field_id"],
                row["field_name"],
                row["field_description"],
                row["value"],
                row["quote"],
            ]
        )
    return BytesIO(buf.getvalue().encode("utf-8-sig"))


def _add_field_line(doc, field_name, value, quote):
    """Single worksheet row: 'Label: answer'."""
    body = Pt(11)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15

    lab = p.add_run(field_name.strip() + ": ")
    _sans_run(lab, size_pt=body, bold=True)

    value_runs = []

    if value:
        lines = value.replace("\r\n", "\n").split("\n")
        vr = p.add_run(lines[0])
        _sans_run(vr, size_pt=body)
        value_runs.append(vr)
        for extra in lines[1:]:
            cont = doc.add_paragraph(extra)
            cont.paragraph_format.space_after = Pt(0)
            cont.paragraph_format.left_indent = Pt(18)
            for r in cont.runs:
                _sans_run(r, size_pt=body)
                value_runs.append(r)
    else:
        ur = p.add_run("_" * 52)
        _sans_run(ur, size_pt=body)
        value_runs.append(ur)

    if quote and value_runs:
        inner = quote.strip().replace('"', "'")
        comment_text = 'Quote from Transcript: "' + inner + '"'
        doc.add_comment(runs=value_runs, text=comment_text, author="")


def _add_choice_line(doc, field_name, value, options, quote):
    """Render a choice/yes_no field with real toggleable Word checkbox SDT controls."""
    body = Pt(11)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15

    lab = p.add_run(field_name.strip() + ":  ")
    _sans_run(lab, size_pt=body, bold=True)

    for i, opt in enumerate(options):
        is_selected = (value == opt)
        _add_checkbox_sdt(p, checked=is_selected)
        spacer = f" {opt}"
        if i < len(options) - 1:
            spacer += "     "
        r = p.add_run(spacer)
        _sans_run(r, size_pt=body)

    if quote:
        inner = quote.strip().replace('"', "'")
        comment_text = 'Quote from Transcript: "' + inner + '"'
        all_runs = list(p.runs)
        if all_runs:
            doc.add_comment(runs=all_runs, text=comment_text, author="")


def _add_section_header(doc, text):
    """Bold section label (e.g. 'Ancillary Document Agents:')."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    _sans_run(r, size_pt=Pt(11), bold=True)


def _add_group_header(doc, text):
    """Bold sub-group label (e.g. 'Trustees for Separate Trusts:')."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    _sans_run(r, size_pt=Pt(11), bold=True)


def _add_section_spacer(doc):
    """Whitespace between sections."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)


def build_questionnaire_docx(sections, schema, repeatable_sections):
    """Worksheet-style .docx (firm drafting-notes layout); returns BytesIO."""
    doc = Document()
    doc.styles["Normal"].font.name = _SANS
    for s in doc.sections:
        s.top_margin = Inches(1)
        s.bottom_margin = Inches(1)
        s.left_margin = Inches(1)
        s.right_margin = Inches(1)

    repeatable_set = set(repeatable_sections or [])
    sections = sections or {}

    for sec_idx, sec in enumerate(schema):
        sid = sec["id"]
        sdata = sections.get(sid) or {}

        if sec_idx > 0:
            _add_section_spacer(doc)

        header = sec.get("header")
        if header:
            _add_section_header(doc, header)

        if sid in repeatable_set:
            entries = sdata.get("entries") or []
            if not entries:
                continue
            for ei, ent in enumerate(entries):
                sub = doc.add_paragraph()
                sub.paragraph_format.space_before = Pt(4)
                sub.paragraph_format.space_after = Pt(4)
                sr = sub.add_run("Entry " + str(ei + 1))
                _sans_run(sr, size_pt=Pt(10), bold=True)

                fmap = ent.get("fields") or {}
                for field in sec.get("fields") or []:
                    _render_docx_field(doc, field, fmap)
        else:
            fmap = sdata.get("fields") or {}
            for field in sec.get("fields") or []:
                _render_docx_field(doc, field, fmap)

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out


def _render_docx_field(doc, field, fmap):
    """Dispatch a single field to the right DOCX renderer based on its type."""
    group_header = field.get("group_header")
    if group_header:
        _add_group_header(doc, group_header)

    fid = field["id"]
    fd = fmap.get(fid) or {}
    val = fd.get("value")
    qv = _str_or_empty(val) if val is not None else ""
    qq = fd.get("quote")
    qq = _str_or_empty(qq) if qq is not None else ""
    field_name = field.get("name") or fid
    field_type = field.get("type", "text")

    condition = field.get("condition")
    if condition:
        cond_fd = fmap.get(condition["field"]) or {}
        cond_val = _str_or_empty(cond_fd.get("value"))
        is_active = (cond_val == condition["value"])
        if not is_active and not qv:
            _add_field_line(doc, field_name, "", qq)
            return

    if field_type == "yes_no":
        _add_choice_line(doc, field_name, qv, ["Yes", "No"], qq)
    elif field_type == "choice":
        _add_choice_line(doc, field_name, qv, field.get("options") or [], qq)
    else:
        _add_field_line(doc, field_name, qv, qq)
